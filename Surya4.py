import io
from typing import List
import pypdfium2
import streamlit as st
from PIL import Image
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

from surya.detection import batch_text_detection
from surya.model.detection.model import load_model, load_processor
from surya.model.recognition.model import load_model as load_rec_model
from surya.model.recognition.processor import load_processor as load_rec_processor
from surya.languages import CODE_TO_LANGUAGE
from surya.ocr import run_ocr
from surya.settings import settings
from surya.postprocessing.text import draw_text_on_image


# Cache models to optimize loading
@st.cache_resource()
def load_det_cached():
    return load_model(), load_processor()


@st.cache_resource()
def load_rec_cached():
    return load_rec_model(), load_rec_processor()


# Function for OCR
def ocr(img, highres_img, langs: List[str]) -> (Image.Image, List[str], List):
    img_pred = run_ocr([img], [langs], det_model, det_processor, rec_model, rec_processor, highres_images=[highres_img])[0]
    bboxes = [l.bbox for l in img_pred.text_lines]
    text_lines = [l.text for l in img_pred.text_lines]
    rec_img = draw_text_on_image(bboxes, text_lines, img.size, langs)
    return rec_img, text_lines, bboxes


# Save OCR output to DOCX
def save_ocr_to_docx(text_lines: List[str], filename: str):
    doc = Document()
    doc.add_heading("OCR Output", level=1)
    for line in text_lines:
        doc.add_paragraph(line)
    doc.save(filename)
    return filename


# Save OCR output to PDF
def save_ocr_to_pdf(image_size, bboxes, text_lines, pdf_filename: str):
    pdf = canvas.Canvas(pdf_filename, pagesize=letter)

    img_width, img_height = image_size
    pdf_width, pdf_height = letter

    for bbox, text in zip(bboxes, text_lines):
        x0, y0, x1, y1 = bbox
        norm_x = x0 / img_width
        norm_y = y0 / img_height
        pdf_x = norm_x * pdf_width
        pdf_y = (1 - norm_y) * pdf_height  # Flip Y-axis
        pdf.drawString(pdf_x, pdf_y, text)

    pdf.save()
    return pdf_filename


# Open PDF and extract page images
def get_all_page_images(pdf_file, dpi=settings.IMAGE_DPI):
    stream = io.BytesIO(pdf_file.getvalue())
    doc = pypdfium2.PdfDocument(stream)
    page_count = len(doc)
    images = []

    for page_num in range(page_count):
        renderer = doc.render(
            pypdfium2.PdfBitmap.to_pil,
            page_indices=[page_num],
            scale=dpi / 72,
        )
        image = list(renderer)[0].convert("RGB")
        images.append(image)
    return images


# Streamlit app UI
st.set_page_config(layout="wide")
st.markdown("""
# OCR Demo with DOCX and PDF Export

This app performs OCR for all pages in a PDF and exports results to DOCX and PDF files.

Notes:
- Works best with printed text documents.
- You can upload PDF or image files (PNG, JPG).
""")

in_file = st.sidebar.file_uploader("PDF file or image:", type=["pdf", "png", "jpg", "jpeg", "gif", "webp"])
languages = st.sidebar.multiselect(
    "Languages", sorted(list(CODE_TO_LANGUAGE.values())),
    default=[], max_selections=4, help="Select known languages in the document to improve OCR accuracy."
)

if in_file is None:
    st.stop()

filetype = in_file.type

# Process all pages of a PDF or single image
if "pdf" in filetype:
    pil_images = get_all_page_images(in_file)
else:
    pil_images = [Image.open(in_file).convert("RGB")]

# Load models
det_model, det_processor = load_det_cached()
rec_model, rec_processor = load_rec_cached()

# OCR processing
if st.sidebar.button("Run OCR"):
    all_text_lines = []
    all_bboxes = []

    for page_index, pil_image in enumerate(pil_images):
        st.markdown(f"### Page {page_index + 1}")
        pil_image_highres = pil_image
        rec_img, text_lines, bboxes = ocr(pil_image, pil_image_highres, languages)

        # Display OCR output for each page
        st.image(rec_img, caption=f"OCR Result (Page {page_index + 1})", use_column_width=True)
        st.text("\n".join(text_lines))
        all_text_lines.extend(text_lines)
        all_bboxes.append((pil_image.size, bboxes, text_lines))

    # Save to DOCX
    if st.sidebar.button("Export to DOCX"):
        docx_filename = "ocr_output.docx"
        save_ocr_to_docx(all_text_lines, docx_filename)
        with open(docx_filename, "rb") as docx_file:
            st.sidebar.download_button(
                label="Download DOCX",
                data=docx_file,
                file_name=docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # Save to PDF
    if st.sidebar.button("Export to PDF"):
        pdf_filename = "ocr_output.pdf"
        pdf = canvas.Canvas(pdf_filename, pagesize=letter)

        for img_size, bboxes, text_lines in all_bboxes:
            img_width, img_height = img_size
            pdf_width, pdf_height = letter

            for bbox, text in zip(bboxes, text_lines):
                x0, y0, x1, y1 = bbox
                norm_x = x0 / img_width
                norm_y = y0 / img_height
                pdf_x = norm_x * pdf_width
                pdf_y = (1 - norm_y) * pdf_height  # Flip Y-axis
                pdf.drawString(pdf_x, pdf_y, text)
            pdf.showPage()

        pdf.save()
        with open(pdf_filename, "rb") as pdf_file:
            st.sidebar.download_button(
                label="Download PDF",
                data=pdf_file,
                file_name=pdf_filename,
                mime="application/pdf"
            )
