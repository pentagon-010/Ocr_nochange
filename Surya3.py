import io
from typing import List
from PIL import Image
import pypdfium2
import streamlit as st
from docx import Document

from surya.ocr import run_ocr
from surya.input.langs import replace_lang_with_code
from surya.postprocessing.text import draw_text_on_image
from surya.model.detection.model import load_model as load_det_model
from surya.model.detection.processor import load_processor as load_det_processor
from surya.model.recognition.model import load_model as load_rec_model
from surya.model.recognition.processor import load_processor as load_rec_processor

# Load models and processors (cached for efficiency)
@st.cache_resource()
def load_ocr_models():
    det_model = load_det_model()
    det_processor = load_det_processor()
    rec_model = load_rec_model()
    rec_processor = load_rec_processor()
    return det_model, det_processor, rec_model, rec_processor

# Extract text and bounding boxes using OCR
def perform_ocr(img, highres_img, langs: List[str]):
    replace_lang_with_code(langs)
    det_model, det_processor, rec_model, rec_processor = load_ocr_models()
    result = run_ocr([img], [langs], det_model, det_processor, rec_model, rec_processor, highres_images=[highres_img])[0]
    return result

# Open PDF and get page as image
def get_pdf_page_image(pdf_file, page_num, dpi):
    doc = pypdfium2.PdfDocument(io.BytesIO(pdf_file.getvalue()))
    renderer = doc.render(
        pypdfium2.PdfBitmap.to_pil,
        page_indices=[page_num - 1],
        scale=dpi / 72,
    )
    return list(renderer)[0].convert("RGB")

# Main Streamlit app
st.set_page_config(layout="wide")
st.title("OCR Text to DOCX")

# File uploader
uploaded_file = st.file_uploader("Upload a PDF or Image:", type=["pdf", "png", "jpg", "jpeg"])
languages = st.sidebar.multiselect("Languages", ["en", "fr", "es", "de"], default=["en"], max_selections=4)

if uploaded_file:
    file_type = uploaded_file.type
    if "pdf" in file_type:
        page_count = pypdfium2.PdfDocument(io.BytesIO(uploaded_file.getvalue())).get_page_count()
        page_num = st.sidebar.number_input("Select Page:", 1, page_count, 1)
        image = get_pdf_page_image(uploaded_file, page_num, dpi=300)
    else:
        image = Image.open(uploaded_file).convert("RGB")
    
    st.image(image, caption="Uploaded Image", use_column_width=True)
    
    if st.button("Run OCR"):
        ocr_result = perform_ocr(image, image, languages)
        text_lines = [(line.text, line.bbox) for line in ocr_result.text_lines]

        # Output OCR data into DOCX
        doc = Document()
        doc.add_heading("OCR Extracted Data", level=1)
        for text, bbox in text_lines:
            doc.add_paragraph(f"Text: {text}")
            doc.add_paragraph(f"Bounding Box: {bbox}")
            doc.add_paragraph("-" * 50)

        doc_path = "ocr_output.docx"
        doc.save(doc_path)

        with open(doc_path, "rb") as f:
            st.download_button("Download DOCX", f, file_name="ocr_output.docx")
