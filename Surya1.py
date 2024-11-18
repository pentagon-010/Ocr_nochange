import io
from typing import List
import pypdfium2
import streamlit as st
from pypdfium2 import PdfiumError
from PIL import Image
from docx import Document  # Import python-docx for DOCX support

from surya.detection import batch_text_detection
from surya.input.pdflines import get_page_text_lines
from surya.model.detection.model import load_model, load_processor
from surya.model.recognition.model import load_model as load_rec_model
from surya.model.recognition.processor import load_processor as load_rec_processor
from surya.languages import CODE_TO_LANGUAGE
from surya.ocr import run_ocr
from surya.settings import settings
from surya.postprocessing.text import draw_text_on_image


# Cache models to optimize loading
@st.cache_resource()
def load_det_cached():
    return load_model(), load_processor()


@st.cache_resource()
def load_rec_cached():
    return load_rec_model(), load_rec_processor()


# Function for OCR
def ocr(img, highres_img, langs: List[str]) -> (Image.Image, List[str]):
    img_pred = run_ocr([img], [langs], det_model, det_processor, rec_model, rec_processor, highres_images=[highres_img])[0]
    bboxes = [l.bbox for l in img_pred.text_lines]
    text_lines = [l.text for l in img_pred.text_lines]
    rec_img = draw_text_on_image(bboxes, text_lines, img.size, langs)
    return rec_img, text_lines


# Save OCR output to DOCX
def save_ocr_to_docx(text_lines: List[str], filename: str):
    doc = Document()
    doc.add_heading("OCR Output", level=1)
    for line in text_lines:
        doc.add_paragraph(line)
    doc.save(filename)
    return filename


# Open PDF and extract page image
def open_pdf(pdf_file):
    stream = io.BytesIO(pdf_file.getvalue())
    return pypdfium2.PdfDocument(stream)


@st.cache_data()
def get_page_image(pdf_file, page_num, dpi=settings.IMAGE_DPI):
    doc = open_pdf(pdf_file)
    renderer = doc.render(
        pypdfium2.PdfBitmap.to_pil,
        page_indices=[page_num - 1],
        scale=dpi / 72,
    )
    png = list(renderer)[0]
    png_image = png.convert("RGB")
    return png_image


# Streamlit app UI
st.set_page_config(layout="wide")
st.markdown("""
# OCR Demo with DOCX Export

This app performs OCR and exports results to a DOCX file.

Notes:
- Works best with printed text documents.
- You can upload PDF or image files (PNG, JPG).
""")

in_file = st.sidebar.file_uploader("PDF file or image:", type=["pdf", "png", "jpg", "jpeg", "gif", "webp"])
languages = st.sidebar.multiselect(
    "Languages", sorted(list(CODE_TO_LANGUAGE.values())),
    default=[], max_selections=4, help="Select known languages in the document to improve OCR accuracy."
)

if in_file is None:
    st.stop()

filetype = in_file.type
if "pdf" in filetype:
    page_number = st.sidebar.number_input("Page number:", min_value=1, value=1)
    pil_image = get_page_image(in_file, page_number)
    pil_image_highres = get_page_image(in_file, page_number, dpi=settings.IMAGE_DPI_HIGHRES)
else:
    pil_image = Image.open(in_file).convert("RGB")
    pil_image_highres = pil_image

# Load models
det_model, det_processor = load_det_cached()
rec_model, rec_processor = load_rec_cached()

# OCR processing
if st.sidebar.button("Run OCR"):
    rec_img, text_lines = ocr(pil_image, pil_image_highres, languages)

    # Display OCR output
    st.image(rec_img, caption="OCR Result", use_column_width=True)
    st.text("\n".join(text_lines))

    # Save to DOCX
    if st.sidebar.button("Export to DOCX"):
        docx_filename = "ocr_output.docx"
        save_ocr_to_docx(text_lines, docx_filename)
        with open(docx_filename, "rb") as docx_file:
            st.sidebar.download_button(
                label="Download DOCX",
                data=docx_file,
                file_name=docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
